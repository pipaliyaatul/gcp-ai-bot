import logging
import os
import time
from typing import Optional, Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json

from config import Config

logger = logging.getLogger(__name__)

class RFPGenerator:
    """Generates RFP summary documents based on extracted content using Vertex AI LLM"""
    
    def __init__(self):
        self.use_ai = os.getenv("USE_AI_FOR_RFP", "true").lower() == "true"
        self.vertex_ai_initialized = False
        self.llm_model = None
        self.model_name = None
        
        # Initialize Vertex AI if project ID is available
        if Config.VERTEX_AI_PROJECT_ID:
            try:
                import vertexai
                from vertexai.generative_models import GenerativeModel
                
                vertexai.init(project=Config.VERTEX_AI_PROJECT_ID, location=Config.VERTEX_AI_LOCATION)
                
                # Try to initialize with the configured model, with fallbacks
                fallback_models = Config.get_fallback_models()
                models_to_try = [Config.VERTEX_AI_MODEL_NAME] + [
                    m for m in fallback_models 
                    if m != Config.VERTEX_AI_MODEL_NAME
                ]
                
                for model_name in models_to_try:
                    try:
                        # Just create the model object - this doesn't make an API call
                        # The actual API call happens when generate_content is called
                        test_model = GenerativeModel(model_name)
                        self.llm_model = test_model
                        self.model_name = model_name
                        self.vertex_ai_initialized = True
                        # Only log once to avoid spam during hot-reload
                        if not hasattr(RFPGenerator, '_llm_logged'):
                            logger.info(f"Vertex AI LLM initialized with model {model_name} for project {Config.VERTEX_AI_PROJECT_ID} in {Config.VERTEX_AI_LOCATION}")
                            RFPGenerator._llm_logged = True
                        break
                    except Exception as model_error:
                        error_msg = str(model_error).lower()
                        if "not found" in error_msg or "404" in error_msg or "does not have access" in error_msg or "publisher model" in error_msg:
                            logger.warning(f"Model {model_name} not available: {model_error}. Trying next model...")
                            continue
                        else:
                            # For other errors during initialization, still try to use it
                            # The error might be during actual generation, not initialization
                            logger.warning(f"Warning with model {model_name} during init: {model_error}. Will try to use it anyway.")
                            self.llm_model = GenerativeModel(model_name)
                            self.model_name = model_name
                            self.vertex_ai_initialized = True
                            break
                
                if not self.vertex_ai_initialized:
                    raise ValueError(f"Could not initialize any Vertex AI model. Tried: {models_to_try}")
                    
            except Exception as e:
                logger.warning(f"Vertex AI LLM initialization failed: {e}. Will use fallback methods.")
                self.vertex_ai_initialized = False
    
    async def generate_rfp_summary(
        self, 
        extracted_text: str, 
        base_template_path: Optional[str] = None,
        base_document_structure: Optional[Dict] = None,
        progress_callback: Optional[callable] = None
    ):
        """
        Generate RFP summary document from extracted text using Vertex AI LLM
        
        Args:
            extracted_text: The extracted text from uploaded file
            base_template_path: Optional path to a base template document
            base_document_structure: Optional structure dict from base document
            progress_callback: Optional callback function(progress: int, message: str) for progress updates
            
        Returns:
            Tuple of (Document, statistics_dict) where statistics_dict contains:
            - model_name: Name of the LLM model used
            - total_tokens: Total tokens used
            - input_tokens: Input tokens
            - output_tokens: Output tokens
            - latency_ms: Latency in milliseconds
            - generation_steps: List of generation steps with details
        """
        statistics = {
            "model_name": self.model_name if self.vertex_ai_initialized else "fallback",
            "total_tokens": 0,
            "input_tokens": 0,
            "output_tokens": 0,
            "latency_ms": 0,
            "generation_steps": [],
            "timestamp": datetime.utcnow().isoformat()
        }
        
        start_time = time.time()
        
        try:
            # Load base template if provided
            base_doc = None
            if base_template_path and os.path.exists(base_template_path):
                try:
                    base_doc = Document(base_template_path)
                    logger.info(f"Loaded base template from {base_template_path}")
                except Exception as e:
                    logger.warning(f"Could not load base template: {e}. Creating new document.")
            
            # Create a new Document (or use template)
            doc = base_doc if base_doc else Document()
            
            # If using Vertex AI LLM, generate content with AI
            if self.vertex_ai_initialized and self.use_ai:
                if base_document_structure:
                    # Align with base document structure
                    doc, step_stats = await self._generate_aligned_with_base(
                        doc, extracted_text, base_document_structure, progress_callback
                    )
                else:
                    doc, step_stats = await self._generate_with_vertex_ai(doc, extracted_text, base_doc is not None, progress_callback)
                statistics["generation_steps"] = step_stats
            else:
                # Fallback to rule-based generation
                doc = await self._generate_fallback(doc, extracted_text)
                statistics["generation_steps"].append({
                    "step": "fallback_generation",
                    "method": "rule_based",
                    "timestamp": datetime.utcnow().isoformat()
                })
            
            # Calculate final statistics - aggregate token usage from all generation steps
            end_time = time.time()
            statistics["latency_ms"] = int((end_time - start_time) * 1000)
            
            # Aggregate token usage from all generation steps
            total_input_tokens = 0
            total_output_tokens = 0
            total_all_tokens = 0
            
            for step in statistics["generation_steps"]:
                if "input_tokens" in step:
                    total_input_tokens += step.get("input_tokens", 0)
                if "output_tokens" in step:
                    total_output_tokens += step.get("output_tokens", 0)
                if "total_tokens" in step:
                    total_all_tokens += step.get("total_tokens", 0)
            
            # Update statistics with aggregated token counts
            statistics["input_tokens"] = total_input_tokens
            statistics["output_tokens"] = total_output_tokens
            statistics["total_tokens"] = total_all_tokens if total_all_tokens > 0 else (total_input_tokens + total_output_tokens)
            
            # Add additional metadata
            statistics["generation_method"] = "single_llm_call" if base_document_structure or len(statistics["generation_steps"]) <= 2 else "multiple_llm_calls"
            statistics["sections_generated"] = len(base_document_structure.get("sections", [])) if base_document_structure else len([
                "Executive Summary", "Introduction", "Background", "Requirements",
                "Technical Specifications", "Timeline", "Budget", "Deliverables", "Conclusion"
            ])
            
            return doc, statistics
            
        except Exception as e:
            logger.error(f"Error generating RFP summary: {str(e)}")
            statistics["error"] = str(e)
            raise
    
    async def _generate_aligned_with_base(
        self,
        doc: Document,
        extracted_text: str,
        base_structure: Dict,
        progress_callback: Optional[callable] = None
    ):
        """Generate RFP document aligned with base document structure"""
        generation_steps = []
        sections = base_structure.get("sections", [])
        total_sections = len(sections)
        
        if progress_callback:
            progress_callback(0, f"Starting alignment with {total_sections} sections...")
        
        # Load base document if available - CRITICAL: Use base document as template
        base_doc = None
        base_file_path = base_structure.get("file_path")
        if base_file_path and os.path.exists(base_file_path):
            try:
                base_doc = Document(base_file_path)
                logger.info(f"Loaded base document from {base_file_path} with {total_sections} sections")
                # Use base document as the starting point - preserves formatting and structure
                doc = base_doc
            except Exception as e:
                logger.warning(f"Could not load base document from {base_file_path}: {e}. Will create new document.")
        else:
            logger.warning(f"Base document file path not found: {base_file_path}. Will create new document.")
        
        # Verify sections were extracted correctly
        if not sections or len(sections) == 0:
            logger.warning("No sections found in base document structure. Using default sections.")
            sections = [
                "Executive Summary",
                "Introduction",
                "Background",
                "Requirements",
                "Technical Specifications",
                "Timeline",
                "Budget",
                "Deliverables",
                "Conclusion"
            ]
            total_sections = len(sections)
        
        # Generate content for all sections in a single LLM call (optimized approach)
        # This is much faster than making separate calls for each section
        if progress_callback:
            progress_callback(20, f"Preparing prompt for all {total_sections} sections...")
        
        section_content_map, token_stats = await self._generate_all_sections_at_once(
            extracted_text,
            sections,
            generation_steps,
            progress_callback
        )
        
        # Store token statistics for aggregation
        generation_steps.append({
            "step": "token_summary",
            "input_tokens": token_stats.get("input_tokens", 0),
            "output_tokens": token_stats.get("output_tokens", 0),
            "total_tokens": token_stats.get("total_tokens", 0),
            "timestamp": datetime.utcnow().isoformat()
        })
        
        if progress_callback:
            progress_callback(85, f"Combining {len(section_content_map)} sections into final document...")
        
        # Replace or add sections in document
        # Find existing sections and replace, or add new ones
        existing_headings = {}
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                existing_headings[paragraph.text.strip()] = paragraph
        
        # Process each section - replace or add content
        total_sections_to_process = len(section_content_map)
        for idx, (section_name, content) in enumerate(section_content_map.items()):
            if progress_callback:
                # Update progress: 85-95% for section combination
                section_progress = 85 + int((idx / total_sections_to_process) * 10)
                progress_callback(section_progress, f"Updating section {idx + 1}/{total_sections_to_process}: {section_name}...")
            # Check if section exists in document
            if section_name in existing_headings:
                # Find the heading and replace content after it
                heading_para = existing_headings[section_name]
                logger.debug(f"Replacing content for existing section: {section_name}")
                
                # Remove paragraphs after heading until next heading
                para = heading_para._element.getnext()
                while para is not None:
                    next_para = para.getnext()
                    if para.tag.endswith('p'):
                        para_text = para.text if hasattr(para, 'text') else ''
                        # Stop if we hit another section heading
                        if para_text and any(para_text.strip().startswith(h) for h in sections if h != section_name):
                            break
                        para.getparent().remove(para)
                    para = next_para
                
                # Add new content
                for line in content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith('#'):
                        # Handle bullet points
                        if line.startswith('-') or line.startswith('*') or line.startswith('•'):
                            bullet_text = line.lstrip('-*•').strip()
                            if bullet_text:
                                doc.add_paragraph(bullet_text, style='List Bullet')
                        else:
                            doc.add_paragraph(line)
            else:
                # Add new section (section not found in base document)
                logger.debug(f"Adding new section: {section_name}")
                doc.add_heading(section_name, level=1)
                for line in content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith('#'):
                        # Handle bullet points
                        if line.startswith('-') or line.startswith('*') or line.startswith('•'):
                            bullet_text = line.lstrip('-*•').strip()
                            if bullet_text:
                                doc.add_paragraph(bullet_text, style='List Bullet')
                        else:
                            doc.add_paragraph(line)
        
        logger.info(f"Successfully processed {len(section_content_map)} sections in base document structure")
        
        if progress_callback:
            progress_callback(95, f"Document structure updated successfully! All {len(section_content_map)} sections have been combined into the final document.")
        
        return doc, generation_steps
    
    async def _generate_all_sections_at_once(
        self,
        extracted_text: str,
        sections: List[str],
        generation_steps: list,
        progress_callback: Optional[callable] = None
    ) -> tuple[Dict[str, str], Dict[str, Any]]:
        """
        Generate content for all sections in a single LLM call using structured JSON output.
        This is much faster than making separate calls for each section.
        
        Returns:
            Dict mapping section names to their content
        """
        try:
            # Prepare comprehensive prompt with all sections
            sections_list = "\n".join([f"{i+1}. {section}" for i, section in enumerate(sections)])
            
            # Use more of the extracted text if available (up to 30K chars for better context)
            text_sample = extracted_text[:30000] if len(extracted_text) > 30000 else extracted_text
            
            combined_prompt = f"""You are an expert RFP document generator. Based on the following content, extract or generate information for ALL the sections listed below.

CONTENT TO ANALYZE:
{text_sample}

SECTIONS TO GENERATE:
{sections_list}

INSTRUCTIONS:
1. For each section, extract relevant information from the content if available
2. If information is not available for a section, use: "[PLACEHOLDER: Information not available from provided content]"
3. Provide clear, well-structured content for each section
4. Return your response as a JSON object with section names as keys and content as values

REQUIRED OUTPUT FORMAT (JSON):
{{
  "Section Name 1": "Content for section 1...",
  "Section Name 2": "Content for section 2...",
  ...
}}

IMPORTANT: Return ONLY valid JSON. Do not include any markdown formatting, code blocks, or explanatory text outside the JSON object."""

            if progress_callback:
                progress_callback(30, f"Calling AI model to generate all {len(sections)} sections at once...")
            
            # Make single LLM call with JSON response format
            response_text, token_stats = await self._call_llm_with_tracing_json(
                combined_prompt,
                "all_sections_batch",
                generation_steps,
                max_output_tokens=8192  # Increased for multiple sections
            )
            
            if progress_callback:
                progress_callback(70, "Parsing AI response and extracting sections...")
            
            # Parse JSON response
            section_content_map = self._parse_sections_json(response_text, sections)
            
            # Return section map and token statistics
            return section_content_map, token_stats
            
            if progress_callback:
                progress_callback(85, f"Successfully generated {len(section_content_map)} sections")
            
            return section_content_map, token_stats
            
        except Exception as e:
            logger.warning(f"Batch generation failed: {str(e)}. Falling back to per-section generation.")
            # Fallback to per-section generation if batch fails
            section_map, token_stats = await self._generate_sections_one_by_one(
                extracted_text,
                sections,
                generation_steps,
                progress_callback
            )
            return section_map, token_stats
    
    async def _generate_sections_one_by_one(
        self,
        extracted_text: str,
        sections: List[str],
        generation_steps: list,
        progress_callback: Optional[callable] = None
    ) -> tuple[Dict[str, str], Dict[str, Any]]:
        """Fallback: Generate content for each section individually (original approach)"""
        section_content_map = {}
        total_sections = len(sections)
        
        for idx, section_name in enumerate(sections):
            if progress_callback:
                progress = int((idx / total_sections) * 80) + 10  # 10-90%
                progress_callback(progress, f"Processing section {idx + 1}/{total_sections}: {section_name}...")
            
            # Generate content for this section
            section_prompt = f"""Based on the following content, extract or generate information for the section "{section_name}".
            If the information is available in the content, extract it. If not available, use "[PLACEHOLDER: Information not available from provided content]".
            Provide clear, relevant content for this section.

            Content:
            {extracted_text[:8000]}

            Section: {section_name}
            Content:"""
            
            section_content, step_token_stats = await self._call_llm_with_tracing(
                section_prompt,
                f"section_{section_name.lower().replace(' ', '_')}",
                generation_steps
            )
            
            section_content_map[section_name] = section_content
        
        # Aggregate token statistics from all steps
        total_input_tokens = sum(step.get("input_tokens", 0) for step in generation_steps)
        total_output_tokens = sum(step.get("output_tokens", 0) for step in generation_steps)
        total_tokens = sum(step.get("total_tokens", 0) for step in generation_steps)
        
        token_stats = {
            "input_tokens": total_input_tokens,
            "output_tokens": total_output_tokens,
            "total_tokens": total_tokens
        }
        
        return section_content_map, token_stats
    
    def _parse_sections_json(self, response_text: str, sections: List[str]) -> Dict[str, str]:
        """Parse JSON response and extract section content"""
        import re
        
        # Try to extract JSON from response (might have markdown code blocks)
        json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
        else:
            json_str = response_text.strip()
        
        # Remove markdown code block markers if present
        json_str = json_str.strip()
        if json_str.startswith('```'):
            json_str = json_str.split('```')[1]
            if json_str.startswith('json'):
                json_str = json_str[4:]
        json_str = json_str.strip()
        
        try:
            parsed = json.loads(json_str)
            section_content_map = {}
            
            # Map parsed JSON to sections (handle case-insensitive matching)
            for section_name in sections:
                # Try exact match first
                if section_name in parsed:
                    section_content_map[section_name] = str(parsed[section_name])
                else:
                    # Try case-insensitive match
                    found = False
                    for key, value in parsed.items():
                        if key.lower() == section_name.lower():
                            section_content_map[section_name] = str(value)
                            found = True
                            break
                    
                    # If still not found, use placeholder
                    if not found:
                        section_content_map[section_name] = "[PLACEHOLDER: Information not available from provided content]"
            
            return section_content_map
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON response: {e}")
            logger.debug(f"Response text: {response_text[:500]}")
            # Return placeholders for all sections if parsing fails
            return {section: "[PLACEHOLDER: Failed to parse AI response]" for section in sections}
    
    async def _call_llm_with_tracing_json(
        self,
        prompt: str,
        step_name: str,
        generation_steps: list,
        max_output_tokens: int = 8192
    ) -> tuple[str, Dict[str, Any]]:
        """Call Vertex AI LLM with JSON response format and tracing"""
        step_start = time.time()
        
        try:
            fallback_models = Config.get_fallback_models()
            models_to_try = [self.model_name] + [
                m for m in fallback_models 
                if m != self.model_name
            ]
            
            response = None
            last_error = None
            successful_model = None
            
            for model_name in models_to_try:
                try:
                    from vertexai.generative_models import GenerativeModel
                    model = GenerativeModel(model_name)
                    
                    # Request JSON response format
                    response = model.generate_content(
                        prompt,
                        generation_config={
                            "temperature": 0.3,  # Lower temperature for more consistent JSON
                            "top_p": 0.95,
                            "top_k": 40,
                            "max_output_tokens": max_output_tokens,
                            "response_mime_type": "application/json",  # Request JSON format
                        }
                    )
                    
                    if model_name != self.model_name:
                        logger.info(f"Switched to fallback model {model_name} for step {step_name}")
                        self.llm_model = model
                        self.model_name = model_name
                    
                    successful_model = model_name
                    break
                    
                except Exception as model_error:
                    error_msg = str(model_error).lower()
                    if "not found" in error_msg or "404" in error_msg or "does not have access" in error_msg:
                        logger.warning(f"Model {model_name} failed: {model_error}. Trying next model...")
                        last_error = model_error
                        continue
                    else:
                        raise
            
            if response is None:
                if last_error:
                    raise last_error
                else:
                    raise ValueError("Failed to generate content with any available model")
            
            step_end = time.time()
            step_latency = int((step_end - step_start) * 1000)
            
            # Extract response text
            if hasattr(response, 'text'):
                response_text = response.text
            elif hasattr(response, 'candidates') and response.candidates:
                response_text = response.candidates[0].content.parts[0].text
            else:
                response_text = str(response)
            
            # Extract token usage if available
            input_tokens = 0
            output_tokens = 0
            total_tokens = 0
            
            if hasattr(response, 'usage_metadata'):
                input_tokens = getattr(response.usage_metadata, 'prompt_token_count', 0)
                output_tokens = getattr(response.usage_metadata, 'candidates_token_count', 0)
                total_tokens = input_tokens + output_tokens
            elif hasattr(response, 'candidates') and response.candidates:
                # Try to get from candidates
                try:
                    if hasattr(response.candidates[0], 'usage_metadata'):
                        input_tokens = getattr(response.candidates[0].usage_metadata, 'prompt_token_count', 0)
                        output_tokens = getattr(response.candidates[0].usage_metadata, 'candidates_token_count', 0)
                        total_tokens = input_tokens + output_tokens
                except:
                    pass
            
            # Log step statistics with token usage
            generation_steps.append({
                "step": step_name,
                "model": successful_model,
                "latency_ms": step_latency,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": total_tokens,
                "timestamp": datetime.utcnow().isoformat()
            })
            
            return response_text, {
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": total_tokens,
                "model": successful_model,
                "latency_ms": step_latency
            }
            
        except Exception as e:
            logger.error(f"LLM call failed for {step_name}: {str(e)}")
            step_end = time.time()
            step_latency = int((step_end - step_start) * 1000)
            generation_steps.append({
                "step": step_name,
                "error": str(e),
                "latency_ms": step_latency,
                "timestamp": datetime.utcnow().isoformat()
            })
            # Return error message with zero token stats
            return f"[AI generation failed for {step_name}: {str(e)}]", {
                "input_tokens": 0,
                "output_tokens": 0,
                "total_tokens": 0,
                "model": None,
                "latency_ms": step_latency
            }
    
    async def _generate_with_vertex_ai(
        self, 
        doc: Document, 
        extracted_text: str, 
        has_template: bool,
        progress_callback: Optional[callable] = None
    ):
        """Generate RFP document using Vertex AI LLM with single call approach (consistent with base document flow)"""
        generation_steps = []
        
        try:
            # Use default sections if no base document structure provided
            default_sections = [
                "Executive Summary",
                "Introduction",
                "Background",
                "Requirements",
                "Technical Specifications",
                "Timeline",
                "Budget",
                "Deliverables",
                "Conclusion"
            ]
            
            # Use single LLM call approach for consistency (same as base document flow)
            if progress_callback:
                progress_callback(10, f"Preparing prompt for all {len(default_sections)} sections...")
            
            # Generate all sections in one call
            section_content_map, token_stats = await self._generate_all_sections_at_once(
                extracted_text,
                default_sections,
                generation_steps,
                progress_callback
            )
            
            # Store token statistics for aggregation
            generation_steps.append({
                "step": "token_summary",
                "input_tokens": token_stats.get("input_tokens", 0),
                "output_tokens": token_stats.get("output_tokens", 0),
                "total_tokens": token_stats.get("total_tokens", 0),
                "timestamp": datetime.utcnow().isoformat()
            })
            
            if progress_callback:
                progress_callback(85, f"Combining {len(section_content_map)} sections into final document...")
            
            # Add title if document is new
            if not has_template:
                if progress_callback:
                    progress_callback(86, "Adding document title and metadata...")
                title = doc.add_heading('RFP Summary Document', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            if not has_template:
                doc.add_heading('Document Information', level=1)
                doc.add_paragraph(f'Generated from: User uploaded content')
                doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
                doc.add_paragraph(f'Content length: {len(extracted_text)} characters')
            
            # Add all sections to document
            total_sections = len(section_content_map)
            for idx, (section_name, content) in enumerate(section_content_map.items()):
                if progress_callback:
                    # Update progress: 87-95% for section addition
                    section_progress = 87 + int((idx / total_sections) * 8)
                    progress_callback(section_progress, f"Adding section {idx + 1}/{total_sections}: {section_name}...")
                doc.add_heading(section_name, level=1)
                # Parse content and add paragraphs
                for line in content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith('#'):
                        # Check if it's a bullet point
                        if line.startswith('-') or line.startswith('*') or line.startswith('•'):
                            bullet_text = line.lstrip('-*•').strip()
                            if bullet_text:
                                doc.add_paragraph(bullet_text, style='List Bullet')
                        else:
                            doc.add_paragraph(line)
            
            if progress_callback:
                progress_callback(95, f"Document generation complete! All {total_sections} sections have been successfully added to the document.")
            
            return doc, generation_steps
            
        except Exception as e:
            logger.error(f"Error in Vertex AI generation: {str(e)}")
            generation_steps.append({
                "step": "error",
                "error": str(e),
                "timestamp": datetime.utcnow().isoformat()
            })
            raise
    
    async def _call_llm_with_tracing(
        self, 
        prompt: str, 
        step_name: str, 
        generation_steps: list
    ) -> tuple[str, Dict[str, Any]]:
        """Call Vertex AI LLM with tracing and statistics collection"""
        step_start = time.time()
        
        try:
            # Try with current model first, then fallbacks if needed
            fallback_models = Config.get_fallback_models()
            models_to_try = [self.model_name] + [
                m for m in fallback_models 
                if m != self.model_name
            ]
            
            response = None
            last_error = None
            successful_model = None
            
            for model_name in models_to_try:
                try:
                    # Create model instance for this attempt
                    from vertexai.generative_models import GenerativeModel
                    model = GenerativeModel(model_name)
                    
                    response = model.generate_content(
                        prompt,
                        generation_config={
                            "temperature": 0.7,
                            "top_p": 0.95,
                            "top_k": 40,
                            "max_output_tokens": 2048,
                        }
                    )
                    
                    # If successful, update the model being used
                    if model_name != self.model_name:
                        logger.info(f"Switched to fallback model {model_name} for step {step_name}")
                        self.llm_model = model
                        self.model_name = model_name
                    
                    successful_model = model_name
                    break  # Success, exit loop
                    
                except Exception as model_error:
                    error_msg = str(model_error).lower()
                    if "not found" in error_msg or "404" in error_msg or "does not have access" in error_msg or "publisher model" in error_msg:
                        logger.warning(f"Model {model_name} failed for step {step_name}: {model_error}. Trying next model...")
                        last_error = model_error
                        continue
                    else:
                        # For other errors, re-raise (might be a different issue)
                        raise
            
            # If we exhausted all models without success, raise the last error
            if response is None:
                if last_error:
                    raise last_error
                else:
                    raise ValueError("Failed to generate content with any available model")
            
            step_end = time.time()
            step_latency = int((step_end - step_start) * 1000)
            
            # Extract response text
            if hasattr(response, 'text'):
                response_text = response.text
            elif hasattr(response, 'candidates') and response.candidates:
                response_text = response.candidates[0].content.parts[0].text
            else:
                response_text = str(response)
            
            # Extract token usage if available
            input_tokens = 0
            output_tokens = 0
            if hasattr(response, 'usage_metadata'):
                input_tokens = getattr(response.usage_metadata, 'prompt_token_count', 0)
                output_tokens = getattr(response.usage_metadata, 'candidates_token_count', 0)
            elif hasattr(response, 'prompt_feedback'):
                # Try alternative way to get token counts
                pass
            
            # Record step statistics
            step_stats = {
                "step": step_name,
                "latency_ms": step_latency,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": input_tokens + output_tokens,
                "timestamp": datetime.utcnow().isoformat(),
                "response_length": len(response_text)
            }
            generation_steps.append(step_stats)
            
            logger.info(f"LLM step '{step_name}' completed in {step_latency}ms, tokens: {input_tokens + output_tokens}")
            
            # Return response text and token statistics
            return response_text, {
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "total_tokens": input_tokens + output_tokens,
                "model": successful_model,
                "latency_ms": step_latency
            }
            
        except Exception as e:
            step_end = time.time()
            step_latency = int((step_end - step_start) * 1000)
            error_msg = str(e)
            
            generation_steps.append({
                "step": step_name,
                "error": error_msg,
                "latency_ms": step_latency,
                "timestamp": datetime.utcnow().isoformat()
            })
            
            logger.error(f"LLM step '{step_name}' failed: {error_msg}")
            # Return a fallback message with zero token stats
            return f"[AI generation failed for {step_name}: {error_msg}]", {
                "input_tokens": 0,
                "output_tokens": 0,
                "total_tokens": 0,
                "model": None,
                "latency_ms": step_latency
            }
    
    async def _generate_fallback(self, doc: Document, extracted_text: str) -> Document:
        """Fallback method using rule-based generation (original implementation)"""
        # Add title
        title = doc.add_heading('RFP Summary Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        doc.add_heading('Document Information', level=1)
        doc.add_paragraph(f'Generated from: User uploaded content')
        doc.add_paragraph(f'Content length: {len(extracted_text)} characters')
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = self._generate_summary_text(extracted_text)
        doc.add_paragraph(summary_text)
        
        # Add key requirements section
        doc.add_heading('Key Requirements', level=1)
        requirements = self._extract_requirements(extracted_text)
        for req in requirements:
            doc.add_paragraph(req, style='List Bullet')
        
        # Add technical specifications
        doc.add_heading('Technical Specifications', level=1)
        tech_specs = self._extract_technical_specs(extracted_text)
        for spec in tech_specs:
            doc.add_paragraph(spec, style='List Bullet')
        
        # Add timeline and milestones
        doc.add_heading('Timeline and Milestones', level=1)
        timeline_text = self._extract_timeline_info(extracted_text)
        doc.add_paragraph(timeline_text if timeline_text else "Timeline information to be determined based on project requirements.")
        
        # Add budget considerations
        doc.add_heading('Budget Considerations', level=1)
        budget_text = self._extract_budget_info(extracted_text)
        doc.add_paragraph(budget_text if budget_text else "Budget information to be discussed during proposal evaluation.")
        
        # Add compliance and standards
        doc.add_heading('Compliance and Standards', level=1)
        compliance_text = self._generate_compliance_section(extracted_text)
        doc.add_paragraph(compliance_text)
        
        # Add next steps
        doc.add_heading('Recommended Next Steps', level=1)
        next_steps = [
            "Review and validate all requirements with stakeholders",
            "Prepare detailed technical proposal",
            "Identify resource requirements and team composition",
            "Develop project timeline and milestones",
            "Prepare cost estimate and budget breakdown",
            "Submit proposal by specified deadline"
        ]
        for step in next_steps:
            doc.add_paragraph(step, style='List Number')
        
        return doc
    
    def _generate_summary_text(self, text: str) -> str:
        """Generate executive summary from extracted text"""
        words = text.split()
        word_count = len(words)
        
        summary = f"""
        This document provides a comprehensive summary of the Request for Proposal (RFP) requirements 
        based on the uploaded content. The source material contains approximately {word_count} words 
        of detailed information that has been analyzed to extract key project requirements, 
        technical specifications, and compliance standards.
        
        The following sections provide a structured breakdown of the RFP requirements to assist 
        in preparing a comprehensive proposal response.
        """
        return summary.strip()
    
    def _extract_requirements(self, text: str) -> list:
        """Extract key requirements from text"""
        requirements = []
        text_lower = text.lower()
        
        requirement_keywords = ['must', 'required', 'shall', 'should', 'need', 'requirement']
        sentences = text.split('.')
        
        for sentence in sentences[:20]:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in requirement_keywords):
                if len(sentence.strip()) > 20:
                    requirements.append(sentence.strip())
        
        if not requirements:
            requirements = [
                "All requirements must be clearly documented and validated",
                "Technical specifications must meet industry standards",
                "Compliance with security and data protection regulations is mandatory"
            ]
        
        return requirements[:10]
    
    def _extract_technical_specs(self, text: str) -> list:
        """Extract technical specifications from text"""
        specs = []
        text_lower = text.lower()
        
        tech_keywords = ['api', 'database', 'server', 'cloud', 'security', 'integration', 'platform']
        sentences = text.split('.')
        
        for sentence in sentences[:20]:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in tech_keywords):
                if len(sentence.strip()) > 20:
                    specs.append(sentence.strip())
        
        if not specs:
            specs = [
                "Technical architecture must be scalable and secure",
                "Integration with existing systems must be seamless",
                "Cloud-based solutions preferred"
            ]
        
        return specs[:10]
    
    def _extract_timeline_info(self, text: str) -> Optional[str]:
        """Extract timeline information from text"""
        timeline_keywords = ['deadline', 'timeline', 'schedule', 'milestone', 'delivery', 'due date']
        text_lower = text.lower()
        
        sentences = text.split('.')
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in timeline_keywords):
                if len(sentence.strip()) > 20:
                    return sentence.strip()
        
        return None
    
    def _extract_budget_info(self, text: str) -> Optional[str]:
        """Extract budget information from text"""
        budget_keywords = ['budget', 'cost', 'price', 'funding', 'financial', 'payment']
        text_lower = text.lower()
        
        sentences = text.split('.')
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in budget_keywords):
                if len(sentence.strip()) > 20:
                    return sentence.strip()
        
        return None
    
    def _generate_compliance_section(self, text: str) -> str:
        """Generate compliance and standards section"""
        return """
        All proposals must comply with industry-standard security practices, data protection 
        regulations (including GDPR, HIPAA where applicable), and accessibility standards. 
        The solution must adhere to best practices for software development, testing, and 
        deployment. Documentation and code quality standards must be maintained throughout 
        the project lifecycle.
        """.strip()
    
    async def chat_with_agent(self, message: str) -> str:
        """Chat with AI agent using Vertex AI LLM"""
        if self.vertex_ai_initialized and self.llm_model:
            try:
                response = self.llm_model.generate_content(
                    f"You are an AI assistant specialized in RFP analysis and document generation. "
                    f"User message: {message}",
                    generation_config={
                        "temperature": 0.7,
                        "max_output_tokens": 1024,
                    }
                )
                
                if hasattr(response, 'text'):
                    return response.text
                elif hasattr(response, 'candidates') and response.candidates:
                    return response.candidates[0].content.parts[0].text
                else:
                    return str(response)
            except Exception as e:
                logger.error(f"Chat error with Vertex AI: {e}")
                # Fall through to default response
        
        # Fallback responses
        responses = {
            "hello": "Hello! I'm here to help you with RFP analysis and document generation.",
            "help": "I can help you upload files (PDF, DOCX, TXT, or audio files) and generate RFP summaries. Try uploading a file to get started!",
        }
        
        message_lower = message.lower()
        for key, response in responses.items():
            if key in message_lower:
                return response
        
        return f"I understand you said: '{message}'. I'm an AI assistant specialized in RFP analysis. Upload a file to generate a comprehensive RFP summary document."
